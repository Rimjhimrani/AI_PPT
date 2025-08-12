import streamlit as st
import requests
from PIL import Image
import io
import base64
from python-pptx import Presentation
from python-pptx.util import Inches, Pt
from python-pptx.dml.color import RGBColor
from python-pptx.enum.text import PP_ALIGN
import openai
import google.generativeai as genai
from stability_sdk import client
import stability_sdk.interfaces.gooseai.generation.generation_pb2 as generation
import json
import os
from datetime import datetime
import docx
import pandas as pd

# Page configuration
st.set_page_config(
    page_title="AI PowerPoint Generator",
    page_icon="📊",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS for better styling
st.markdown("""
<style>
    .main-header {
        font-size: 2.5rem;
        color: #1f77b4;
        text-align: center;
        margin-bottom: 2rem;
    }
    .step-header {
        background: linear-gradient(90deg, #667eea 0%, #764ba2 100%);
        color: white;
        padding: 1rem;
        border-radius: 10px;
        margin: 1rem 0;
    }
    .info-box {
        background-color: #f0f2f6;
        padding: 1rem;
        border-radius: 10px;
        border-left: 4px solid #1f77b4;
        margin: 1rem 0;
    }
    .success-box {
        background-color: #d4edda;
        padding: 1rem;
        border-radius: 10px;
        border-left: 4px solid #28a745;
        margin: 1rem 0;
    }
</style>
""", unsafe_allow_html=True)

class PowerPointGenerator:
    def __init__(self):
        self.presentation = Presentation()
        self.slides_content = []
        
    def setup_apis(self):
        """Setup API keys from Streamlit secrets or user input"""
        st.sidebar.header("🔑 API Configuration")
        
        # OpenAI API Key
        openai_key = st.sidebar.text_input(
            "OpenAI API Key", 
            type="password",
            value=st.secrets.get("OPENAI_API_KEY", "")
        )
        
        # Google AI API Key (for Gemini)
        google_key = st.sidebar.text_input(
            "Google AI API Key", 
            type="password",
            value=st.secrets.get("GOOGLE_API_KEY", "")
        )
        
        # Stability AI API Key (for image generation)
        stability_key = st.sidebar.text_input(
            "Stability AI API Key", 
            type="password",
            value=st.secrets.get("STABILITY_API_KEY", "")
        )
        
        if openai_key:
            openai.api_key = openai_key
        if google_key:
            genai.configure(api_key=google_key)
        
        return openai_key, google_key, stability_key
    
    def search_web(self, query):
        """Search web for information using SerpAPI or similar"""
        try:
            # Using a free alternative - you can replace with SerpAPI
            search_url = f"https://api.duckduckgo.com/search"
            params = {
                'q': query,
                'format': 'json',
                'no_html': '1',
                'skip_disambig': '1'
            }
            
            response = requests.get(search_url, params=params)
            if response.status_code == 200:
                data = response.json()
                results = []
                for result in data.get('RelatedTopics', [])[:5]:
                    if 'Text' in result:
                        results.append(result['Text'])
                return results
        except Exception as e:
            st.error(f"Web search failed: {str(e)}")
            return []
    
    def analyze_image(self, image, google_key):
        """Analyze uploaded image using Google Gemini Vision"""
        try:
            if not google_key:
                return "Image analysis requires Google AI API key"
            
            model = genai.GenerativeModel('gemini-pro-vision')
            response = model.generate_content([
                "Describe this image in detail, including any text, objects, people, or relevant information that could be used in a presentation:",
                image
            ])
            return response.text
        except Exception as e:
            return f"Image analysis failed: {str(e)}"
    
    def generate_ai_image(self, prompt, stability_key):
        """Generate AI images using Stability AI"""
        try:
            if not stability_key:
                return None
            
            stability_api = client.StabilityInference(
                key=stability_key,
                verbose=True,
            )
            
            answers = stability_api.generate(
                prompt=prompt,
                seed=992446758,
                steps=30,
                cfg_scale=8.0,
                width=1024,
                height=768,
                samples=1,
                sampler=generation.SAMPLER_K_DPMPP_2M
            )
            
            for resp in answers:
                for artifact in resp.artifacts:
                    if artifact.finish_reason == generation.FILTER:
                        st.warning("Image generation filtered due to content policy")
                        return None
                    if artifact.type == generation.ARTIFACT_IMAGE:
                        img = Image.open(io.BytesIO(artifact.binary))
                        return img
        except Exception as e:
            st.error(f"Image generation failed: {str(e)}")
            return None
    
    def generate_presentation_content(self, topic, research_data, openai_key):
        """Generate presentation content using OpenAI GPT"""
        try:
            if not openai_key:
                return self.generate_basic_content(topic, research_data)
            
            # Prepare research context
            research_context = "\n".join(research_data) if research_data else ""
            
            prompt = f"""
            Create a professional PowerPoint presentation about: {topic}
            
            Research context: {research_context}
            
            Generate 8-10 slides with the following structure:
            1. Title slide
            2. Overview/Agenda
            3. 5-7 content slides
            4. Conclusion
            5. Q&A slide
            
            For each slide, provide:
            - Slide title
            - 3-5 bullet points (concise)
            - Speaker notes (detailed explanation)
            - Image prompt (description for AI image generation)
            
            Format as JSON:
            {{
                "slides": [
                    {{
                        "title": "Slide Title",
                        "bullets": ["Point 1", "Point 2", "Point 3"],
                        "speaker_notes": "Detailed explanation...",
                        "image_prompt": "Description for image generation"
                    }}
                ]
            }}
            """
            
            response = openai.ChatCompletion.create(
                model="gpt-3.5-turbo",
                messages=[{"role": "user", "content": prompt}],
                max_tokens=2000,
                temperature=0.7
            )
            
            content = json.loads(response.choices[0].message.content)
            return content["slides"]
            
        except Exception as e:
            st.error(f"Content generation failed: {str(e)}")
            return self.generate_basic_content(topic, research_data)
    
    def generate_basic_content(self, topic, research_data):
        """Generate basic content when AI is not available"""
        slides = [
            {
                "title": f"{topic}",
                "bullets": ["Welcome", "Overview of today's presentation", "Key objectives"],
                "speaker_notes": f"Welcome everyone to this presentation about {topic}. Today we'll cover the key aspects and insights.",
                "image_prompt": f"Professional title slide background for {topic}"
            },
            {
                "title": "Agenda",
                "bullets": ["Introduction", "Main Topics", "Key Findings", "Conclusion"],
                "speaker_notes": "Here's what we'll cover in today's presentation.",
                "image_prompt": "Clean agenda or roadmap visual"
            },
            {
                "title": f"Introduction to {topic}",
                "bullets": research_data[:3] if research_data else ["Key concept 1", "Key concept 2", "Key concept 3"],
                "speaker_notes": f"Let's start with an introduction to {topic}.",
                "image_prompt": f"Introductory image related to {topic}"
            },
            {
                "title": "Key Points",
                "bullets": research_data[3:6] if len(research_data) > 3 else ["Important point 1", "Important point 2", "Important point 3"],
                "speaker_notes": "These are the key points we need to understand.",
                "image_prompt": f"Visual representation of key concepts in {topic}"
            },
            {
                "title": "Conclusion",
                "bullets": ["Summary of key points", "Next steps", "Thank you"],
                "speaker_notes": "To conclude, let's summarize what we've learned today.",
                "image_prompt": "Professional conclusion slide background"
            }
        ]
        return slides
    
    def create_powerpoint(self, slides_content, generated_images=None):
        """Create PowerPoint presentation"""
        prs = Presentation()
        
        for i, slide_data in enumerate(slides_content):
            # Create slide
            if i == 0:  # Title slide
                slide_layout = prs.slide_layouts[0]
            else:
                slide_layout = prs.slide_layouts[1]
            
            slide = prs.slides.add_slide(slide_layout)
            
            # Add title
            title = slide.shapes.title
            title.text = slide_data["title"]
            
            # Add content
            if slide_layout == prs.slide_layouts[1]:  # Content slide
                content = slide.placeholders[1]
                tf = content.text_frame
                tf.text = slide_data["bullets"][0] if slide_data["bullets"] else ""
                
                for bullet in slide_data["bullets"][1:]:
                    p = tf.add_paragraph()
                    p.text = bullet
                    p.level = 0
            
            # Add image if available
            if generated_images and i < len(generated_images) and generated_images[i]:
                img_stream = io.BytesIO()
                generated_images[i].save(img_stream, format='PNG')
                img_stream.seek(0)
                
                slide.shapes.add_picture(
                    img_stream, 
                    Inches(6), Inches(1), 
                    Inches(3), Inches(3)
                )
            
            # Add speaker notes
            notes_slide = slide.notes_slide
            text_frame = notes_slide.notes_text_frame
            text_frame.text = slide_data["speaker_notes"]
        
        return prs
    
    def read_uploaded_file(self, uploaded_file):
        """Read content from uploaded files"""
        content = ""
        
        if uploaded_file.type == "text/plain":
            content = str(uploaded_file.read(), "utf-8")
        
        elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            doc = docx.Document(uploaded_file)
            for paragraph in doc.paragraphs:
                content += paragraph.text + "\n"
        
        elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.presentationml.presentation":
            # Basic PPTX reading - you might need python-pptx
            content = "PowerPoint content uploaded (basic extraction needed)"
        
        return content

def main():
    st.markdown('<h1 class="main-header">🎯 AI PowerPoint Generator</h1>', unsafe_allow_html=True)
    
    generator = PowerPointGenerator()
    
    # Setup APIs
    openai_key, google_key, stability_key = generator.setup_apis()
    
    # Main interface
    col1, col2 = st.columns([2, 1])
    
    with col1:
        st.markdown('<div class="step-header">📝 Step 1: Input Your Content</div>', unsafe_allow_html=True)
        
        # Input options
        input_method = st.radio(
            "Choose input method:",
            ["Text Input", "File Upload", "Topic Only"]
        )
        
        topic = ""
        uploaded_content = ""
        
        if input_method == "Text Input":
            topic = st.text_input("Presentation Topic:", placeholder="e.g., Climate Change Solutions")
            uploaded_content = st.text_area("Additional Content:", height=150, placeholder="Paste your content here...")
        
        elif input_method == "File Upload":
            topic = st.text_input("Presentation Topic:", placeholder="e.g., Climate Change Solutions")
            uploaded_file = st.file_uploader(
                "Upload file:", 
                type=['txt', 'docx', 'pptx'],
                help="Supported formats: TXT, DOCX, PPTX"
            )
            
            if uploaded_file:
                uploaded_content = generator.read_uploaded_file(uploaded_file)
                st.success("File uploaded successfully!")
        
        else:  # Topic Only
            topic = st.text_input("Presentation Topic:", placeholder="e.g., Climate Change Solutions")
        
        # Image upload for analysis
        st.markdown('<div class="step-header">🖼️ Step 2: Upload Images (Optional)</div>', unsafe_allow_html=True)
        uploaded_images = st.file_uploader(
            "Upload images for analysis:",
            type=['png', 'jpg', 'jpeg'],
            accept_multiple_files=True
        )
        
        image_descriptions = []
        if uploaded_images and google_key:
            for img_file in uploaded_images:
                img = Image.open(img_file)
                description = generator.analyze_image(img, google_key)
                image_descriptions.append(description)
                st.image(img, caption=f"Analysis: {description[:100]}...", width=300)
    
    with col2:
        st.markdown('<div class="info-box">ℹ️ <strong>How it works:</strong><br>1. Enter your topic or upload content<br>2. AI researches the web<br>3. Generates structured slides<br>4. Creates AI images<br>5. Builds PowerPoint file</div>', unsafe_allow_html=True)
        
        # Presentation settings
        st.markdown("⚙️ **Settings**")
        num_slides = st.slider("Number of slides:", 5, 15, 8)
        style = st.selectbox("Presentation style:", ["Professional", "Creative", "Academic", "Minimal"])
        include_images = st.checkbox("Generate AI images", value=True)
    
    # Generate presentation
    if st.button("🚀 Generate Presentation", type="primary"):
        if not topic:
            st.error("Please provide a topic for your presentation!")
            return
        
        progress_bar = st.progress(0)
        status_text = st.empty()
        
        try:
            # Step 1: Web research
            status_text.text("🔍 Researching topic...")
            progress_bar.progress(20)
            research_data = generator.search_web(topic + " " + uploaded_content)
            
            # Step 2: Generate content
            status_text.text("✍️ Generating presentation content...")
            progress_bar.progress(40)
            slides_content = generator.generate_presentation_content(
                topic, research_data + image_descriptions, openai_key
            )
            
            # Step 3: Generate images
            generated_images = []
            if include_images and stability_key:
                status_text.text("🎨 Generating AI images...")
                progress_bar.progress(60)
                
                for slide in slides_content:
                    img = generator.generate_ai_image(slide["image_prompt"], stability_key)
                    generated_images.append(img)
            
            # Step 4: Create PowerPoint
            status_text.text("📊 Creating PowerPoint...")
            progress_bar.progress(80)
            prs = generator.create_powerpoint(slides_content, generated_images)
            
            # Step 5: Save and download
            status_text.text("💾 Preparing download...")
            progress_bar.progress(100)
            
            # Save to bytes
            ppt_bytes = io.BytesIO()
            prs.save(ppt_bytes)
            ppt_bytes.seek(0)
            
            # Download button
            st.download_button(
                label="📥 Download PowerPoint",
                data=ppt_bytes,
                file_name=f"{topic.replace(' ', '_')}_presentation.pptx",
                mime="application/vnd.openxmlformats-officedocument.presentationml.presentation"
            )
            
            status_text.text("✅ Presentation ready!")
            st.success("Your presentation has been generated successfully!")
            
            # Preview slides
            st.markdown("### 👀 Preview")
            for i, slide in enumerate(slides_content):
                with st.expander(f"Slide {i+1}: {slide['title']}"):
                    st.write("**Content:**")
                    for bullet in slide['bullets']:
                        st.write(f"• {bullet}")
                    st.write("**Speaker Notes:**")
                    st.write(slide['speaker_notes'])
                    if generated_images and i < len(generated_images) and generated_images[i]:
                        st.image(generated_images[i], width=400)
            
        except Exception as e:
            st.error(f"An error occurred: {str(e)}")
            status_text.text("❌ Generation failed")

# Sidebar information
st.sidebar.markdown("---")
st.sidebar.markdown("### 📋 Required APIs")
st.sidebar.markdown("""
- **OpenAI**: Content generation
- **Google AI**: Image analysis  
- **Stability AI**: Image generation
- **SerpAPI**: Web search (optional)
""")

st.sidebar.markdown("### 💡 Tips")
st.sidebar.markdown("""
- Be specific with your topic
- Upload relevant documents
- Use high-quality images
- Review generated content
""")

if __name__ == "__main__":
    main()
